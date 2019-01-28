# coding=utf-8
from __future__ import unicode_literals

import json
import logging
import os
from datetime import datetime

from docxtpl import DocxTemplate, InlineImage, Document
from docx.shared import Mm, Inches, Cm
from soc_assets.models import TermGroup

logger = logging.getLogger('soc_ssa')


def make_template_docx(template, path, data):
    """
    生成docx
    """
    # 模版路径
    template_path = ""
    term_group_count = TermGroup.objects.all().count()
    content = dict()
    if template.id in [1, 2, 3, 4, 5, 11]:  # 终端安全检查分析报告
        template_path = os.path.join("soc_ssa", "tools", "make_docx", "template", "terminal.docx")
        doc = DocxTemplate(template_path)
        content['module'] = data[0]['params']['title']
        content['time'] = data[0]['params']['time']
        # 1.1  违规总量占比
        content['img11'] = InlineImage(doc, data[3]['chart_path'], width=Mm(160))
        # 1.2  各单位违规类总量统计
        content['img12'] = InlineImage(doc, data[5]['chart_path'], width=Mm(160))
        # 1.3  违规类型详情统计
        content['tab13'] = data[7]['data']['data']

        data_list = []
        for index, item in enumerate(data):
            if index > 8:
                if item['type'] in ['module', 'title1', 'title2', 'title3', 'title4', 'text']:
                    data_list.append({'type': item['type'], 'title': item['params']['title']})
                else:
                    if item['chart_type'] in ['table', 'table_merge']:  # 表格类型
                        data_list.append({
                            'type': item['type'],
                            'chart_type': item['chart_type'],
                            'data': item['data']['data'],
                            'merges': len(item['merges'])
                        })
                    else:
                        img = 'img{}'.format(index)
                        content[img] = InlineImage(doc, item['chart_path'], width=Mm(160))
                        data_list.append({
                            'type': item['type'],
                            'chart_type': item['chart_type'],
                            'chart_path': img})
                        # data_list.append(item)
        content['data_list'] = data_list
        doc.render({'content': content})
        doc.save(path)

    elif template.id in [6, 7, 8, 9, 10, 12]:  # 内网安全统计报表
        template_path = os.path.join("soc_ssa", "tools", "make_docx", "template", "innernet.docx")
        doc = DocxTemplate(template_path)
        '''
        content['module'] = data[0]['params']['title']
        content['time'] = data[0]['params']['time']
        # 一、内网对接资产设备情况
        content['tab1'] = data[2]['data']['data']
        # 二、各资产设备事件量占比
        content['img2'] = InlineImage(doc, data[4]['chart_path'], width=Mm(160))
        # 三、事件趋势统计
        content['img3'] = InlineImage(doc, data[6]['chart_path'], width=Mm(160))
        # 4.1  风险事件统计
        content['tab41'] = data[9]['data']['data']
        # 4.2  风险资产统计
        content['tab42'] = data[11]['data']['data']
        # 4.3  风险部门统计
        content['tab43'] = data[13]['data']['data']
        # 5.1  违规来源统计
        content['img51'] = InlineImage(doc, data[16]['chart_path'], width=Mm(160))
        # 5.2  违规类型统计
        content['tab52'] = data[18]['data']['data']
        # 5.3  违规终端统计
        content['tab53'] = data[20]['data']['data']
        # 5.4  违规单位统计
        content['tab54'] = data[22]['data']['data']
        # 6.1  病毒类型统计
        content['img61'] = InlineImage(doc, data[25]['chart_path'], width=Mm(160))
        # 6.2  感染病毒终端统计（未处理）
        content['tab62'] = data[27]['data']['data']
        # 7.1  攻击类型统计
        content['img71'] = InlineImage(doc, data[30]['chart_path'], width=Mm(160))
        # 7.2  攻击源统计TOP10
        content['img72'] = InlineImage(doc, data[32]['chart_path'], width=Mm(160))
        # 7.3  被攻击终端统计
        content['tab73'] = data[34]['data']['data']
        # 8.1  各部门开关机详情
        tab81, title81, data81 = [], [], []
        for index in range(term_group_count * 2):
            if index % 2 == 0:
                title81.append(data[37 + index]['params']['title'])
            else:
                data81.append(data[37 + index]['data']['data'])
        for index, tit in enumerate(title81):
            tab81.append({'title': tit, 'data': data81[index]})
        content['tab81'] = tab81
        # 当前索引
        current_index = 37 + 2 * term_group_count
        # 8.2  登录失败统计TOP10
        content['img82'] = InlineImage(doc, data[current_index + 1]['chart_path'], width=Mm(160))
        # 8.3.1  打印情况统计TOP10
        content['img831'] = InlineImage(doc, data[current_index + 4]['chart_path'], width=Mm(160))
        # 8.3.2  各部门打印详情
        tab832, title832, data832 = [], [], []
        for index in range(term_group_count * 2):
            if index % 2 == 0:
                title832.append(data[current_index + 6 + index]['params']['title'])
            else:
                data832.append(data[current_index + 6 + index]['data']['data'])
        for index, tit in enumerate(title832):
            tab832.append({'title': tit, 'data': data832[index]})
        content['tab832'] = tab832
        current_index = current_index + 6 + 2 * term_group_count
        # 8.3.3  下班打印详情
        content['tab833'] = data[current_index + 1]['data']['data']
        # 8.4.1  移动盘使用统计TOP10
        content['img841'] = InlineImage(doc, data[current_index + 4]['chart_path'], width=Mm(160))
        # 8.4.2  移动盘使用统计TOP10
        tab842, title842, data842 = [], [], []
        for index in range(term_group_count * 2):
            if index % 2 == 0:
                title842.append(data[current_index + 6 + index]['params']['title'])
            else:
                data842.append(data[current_index + 6 + index]['data']['data'])
        for index, tit in enumerate(title842):
            tab842.append({'title': tit, 'data': data842[index]})
        content['tab842'] = tab842
        current_index = current_index + 6 + 2 * term_group_count
        # 8.5.1  明文存储终端统计TOP10
        content['img851'] = InlineImage(doc, data[current_index + 2]['chart_path'], width=Mm(160))
        # 8.5.2  各单位明文存储详情
        tab852, title852, data852 = [], [], []
        for index in range(term_group_count * 2):
            if index % 2 == 0:
                title852.append(data[current_index + 4 + index]['params']['title'])
            else:
                data852.append(data[current_index + 4 + index]['data']['data'])
        for index, tit in enumerate(title852):
            tab852.append({'title': tit, 'data': data852[index]})
        content['tab852'] = tab852
        current_index = current_index + 4 + 2 * term_group_count
        # 9.1  终端违规操作数据库统计TOP10
        content['img91'] = InlineImage(doc, data[current_index + 2]['chart_path'], width=Mm(160))
        # 9.2  数据库操作违规类型统计
        content['img92'] = InlineImage(doc, data[current_index + 4]['chart_path'], width=Mm(160))
        # 10.1  防火墙
        content['img1011'] = InlineImage(doc, data[current_index + 7]['chart_path'], width=Mm(160))
        content['img1012'] = InlineImage(doc, data[current_index + 8]['chart_path'], width=Mm(160))
        content['img1013'] = InlineImage(doc, data[current_index + 9]['chart_path'], width=Mm(160))
        content['img1014'] = InlineImage(doc, data[current_index + 10]['chart_path'], width=Mm(160))
        content['img1015'] = InlineImage(doc, data[current_index + 11]['chart_path'], width=Mm(160))
        # 10.2  网闸
        content['img1021'] = InlineImage(doc, data[current_index + 13]['chart_path'], width=Mm(160))
        content['img1022'] = InlineImage(doc, data[current_index + 14]['chart_path'], width=Mm(160))
        content['img1023'] = InlineImage(doc, data[current_index + 15]['chart_path'], width=Mm(160))
        content['img1024'] = InlineImage(doc, data[current_index + 16]['chart_path'], width=Mm(160))
        # 十一、入侵监测事件统计
        content['img1101'] = InlineImage(doc, data[current_index + 18]['chart_path'], width=Mm(160))
        content['img1102'] = InlineImage(doc, data[current_index + 19]['chart_path'], width=Mm(160))
        content['img1103'] = InlineImage(doc, data[current_index + 20]['chart_path'], width=Mm(160))
        content['img1104'] = InlineImage(doc, data[current_index + 21]['chart_path'], width=Mm(160))
        '''

        # 测试 合并文档
        content['sub'] = create_sub_docx(doc=doc,
                                         sub_docx_name=datetime.now().strftime("%Y%m%d%H%M%S"),
                                         table_data=data[2]['data'],
                                         merges=data[2]['merges'],
                                         widths=[])

        # doc.render({'content': content})
        # doc.save(path)

    else:
        pass
    if template_path:
        for index, item in enumerate(data):
            if item['type'] in ['module', 'title1', 'title2', 'title3', 'title4', 'text']:
                # print index, item['type'], item['params']['title']
                pass
            else:
                if item['chart_type'] in ['table', 'table_merge']:  # 表格类型
                    print index, item['type'], item['chart_type'], '##############'
                else:
                    pass
                    # print index, item['type'], item['chart_type'], '----------------------', item['chart_path']


def create_sub_docx(doc, sub_docx_name, table_data, merges, widths):
    '''
    生成子文档
    '''
    path = os.path.join('media', 'reports', 'tmp_docx', '{}.docx'.format(sub_docx_name))
    sub_path = add_table_merge(path=path, data=table_data, merges=merges, widths=widths)
    sub = doc.new_subdoc()
    sub.subdocx = Document(sub_path)
    return sub


def add_table_merge(path, data, merges, widths=[]):
    '''
    创建表格
    '''
    logger.info('--生成合并表格{}'.format(datetime.now()))
    labels = data['labels']
    rows = data['data']
    docx = Document()
    style = docx.styles['Table Grid']
    table = docx.add_table(rows=1 + len(rows), cols=len(labels), style=style)
    # 自定义表格宽度
    if widths:
        table.autofit = False  # 很重要！
        for row in range(1 + len(rows)):
            for index, w in enumerate(widths):
                table.cell(row, index).width = Inches(float(w))
    # 生成表头
    hdr_cells = table.rows[0].cells
    for index, value in enumerate(labels):
        if not isinstance(value, unicode):
            value = unicode(value, "utf-8")
        hdr_cells[index].text = value
    format_data = formate_table_merge_data(data=rows, merges=merges)
    logger.info('--格式化表格数据{}'.format(datetime.now()))
    # 生成表格内容
    create_docx_table_body(docx=docx, path=path, table=table, data=format_data, merges=merges)
    logger.info('--合并表格完成{}'.format(datetime.now()))
    docx.save(path)
    return path


def formate_table_merge_data(data, merges):
    '''
    封装表格合并数据
    '''
    name_list = []
    index_list = []
    for index, item in enumerate(data):
        if index != 0 and name_list[-1] == item[0]:
            index_list[-1].append(index)
        else:
            name_list.append(item[0])
            index_list.append([index])
    indexs = []
    for item in index_list:
        if len(item) > 1:
            indexs.append([item[0], item[-1]])
            for index in item:
                if index != item[0]:
                    for m in merges:
                        data[index][m] = ''

    result = {
        'index': indexs,
        'data': data
    }
    return result


def create_docx_table_body(docx, path, table, data, merges):
    '''
    创建表格内容
    '''
    for row_index, item in enumerate(data['data']):
        row_cells = table.rows[row_index + 1].cells
        for index, value in enumerate(item):
            if isinstance(value, int) or isinstance(value, float):
                value = str(value)
            if not isinstance(value, unicode):
                value = unicode(value, "utf-8")
            row_cells[index].text = value
    # logger.info('--开始合并单元格{}'.format(datetime.now()))
    for index in data['index']:
        for m in merges:
            cell_start = table.cell(index[0] + 1, m)
            cell_end = table.cell(index[1] + 1, m)
            cell_start.merge(cell_end)
    # logger.info('--合并单元格结束{}'.format(datetime.now()))
    docx.save(path)
