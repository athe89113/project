# coding=utf-8
from __future__ import unicode_literals

import copy
import logging
from datetime import datetime

try:
    from docx import Document
    from docx.shared import Inches
    from docx.oxml.ns import qn
    from docx.enum.style import WD_STYLE_TYPE
    from docx.shared import Pt
    from docx.enum.text import WD_ALIGN_PARAGRAPH
except ImportError:
    print("Please Install docx")

logger = logging.getLogger('soc_ssa')
from soc_ssa.models import SSAReportTemplate
from soc_ssa.tools.make_docx import docx_freemarker


class MakeDocx(object):
    """
    通过ES中数据生成图表
    """

    def __init__(self, template_id):
        self.template_id = template_id
        self.docx = Document()
        self.docx.styles['Normal'].font.name = u'宋体'
        self.docx.styles['Normal']._element.rPr.rFonts.set(qn('w:eastAsia'), u'宋体')
        self.index = 1
        self.styles = self.docx.styles

    def data(self):
        """
        """

        data = [
            {"time": "2018-04-10 12:12:12", "name": "f1", "value": 1},
            {"time": "2018-04-10 12:12:12", "name": "f2", "value": 2},
            {"time": "2018-04-10 12:12:12", "name": "f3", "value": 3}
        ]
        return data

    def generate_section(self, title, desc, remark, chart=None, table_data=None, data_type=None):
        """
        生成段落
        """
        paragraph = self.docx.add_paragraph(desc)
        run = paragraph.add_run(text=desc)
        run.font.name = u'宋体'
        run.font.size = Pt(10)

        if chart:
            self.docx.add_picture(chart, width=Inches(6.2))
        if table_data:
            self.add_table(name='', data=table_data, data_type=data_type)

        self.docx.add_paragraph(remark)

        self.index += 1

    def create_chart_image(self, path):
        '''
        创建图片
        '''
        self.docx.add_picture(path, width=Inches(6.2))

    def add_table(self, name, data, data_type):
        """
        添加表格
        """
        if name:
            self.create_docx_text(content=name)
        labels = [''] + data['labels']
        rows = []
        for i in data['data']:
            if data_type == 99:
                rows.append([i['name']] + i['data'])
            else:
                rows.append([data['data'].index(i) + 1] + i['data'])
        style = self.docx.styles['Table Grid']
        table = self.docx.add_table(rows=1, cols=len(labels), style=style)
        hdr_cells = table.rows[0].cells
        for index, value in enumerate(labels):
            hdr_cells[index].text = value
        for item in rows:
            row_cells = table.add_row().cells
            for index, value in enumerate(item):
                if isinstance(value, int) or isinstance(value, float):
                    value = str(value)
                if not isinstance(value, unicode):
                    value = unicode(value, "utf-8")
                row_cells[index].text = value

    def add_table_merge(self, path, data, merges, widths=[]):
        '''
        创建表格
        '''
        # logger.info('--生成合并表格{}'.format(datetime.now()))
        labels = data['labels']
        rows = data['data']
        style = self.docx.styles['Table Grid']
        table = self.docx.add_table(rows=1 + len(rows), cols=len(labels), style=style)
        # 自定义表格宽度
        if widths:
            table.autofit = False  # 很重要！
            for row in range(1 + len(rows)):
                for index, w in enumerate(widths):
                    table.cell(row, index).width = Inches(float(w))
        # 生成表头
        hdr_cells = table.rows[0].cells
        for index, value in enumerate(labels):
            if not isinstance(value, unicode):
                value = unicode(value, "utf-8")
            hdr_cells[index].text = value
        if len(merges) > 0:
            format_data = self.formate_table_merge_data(data=rows, merges=merges)
        else:
            format_data = {'data': rows}
        # logger.info('--格式化表格数据{}'.format(datetime.now()))
        # 生成表格内容
        self.create_docx_table_body(path=path, table=table, data=format_data, merges=merges)
        # logger.info('--合并表格完成{}'.format(datetime.now()))

    def formate_table_merge_data(self, data, merges):
        '''
        封装表格合并数据
        '''
        name_list = []
        index_list = []
        for index, item in enumerate(data):
            if index != 0 and name_list[-1] == item[0]:
                index_list[-1].append(index)
            else:
                name_list.append(item[0])
                index_list.append([index])
        indexs = []
        for item in index_list:
            if len(item) > 1:
                indexs.append([item[0], item[-1]])
                for index in item:
                    if index != item[0]:
                        for m in merges:
                            data[index][m] = ''

        result = {
            'index': indexs,
            'data': data
        }
        return result

    def create_docx_table_body(self, path, table, data, merges):
        '''
        创建表格内容
        '''
        for row_index, item in enumerate(data['data']):
            row_cells = table.rows[row_index + 1].cells
            for index, value in enumerate(item):
                if isinstance(value, int) or isinstance(value, float):
                    value = str(value)
                if not isinstance(value, unicode):
                    value = unicode(value, "utf-8")
                row_cells[index].text = value
        if len(merges) > 0:
            # logger.info('--开始合并单元格{}'.format(datetime.now()))
            for index in data['index']:
                for m in merges:
                    cell_start = table.cell(index[0] + 1, m)
                    cell_end = table.cell(index[1] + 1, m)
                    cell_start.merge(cell_end)
                    # logger.info('--合并单元格结束{}'.format(datetime.now()))
                    # self.docx.save(path)

    def generate_docx(self, path, data):
        """
        生成docx
        version 为日期版本
        """
        if not data:
            return ValueError("模版不存在")
        template_obj = SSAReportTemplate.objects.get(id=self.template_id)

        docx_freemarker.make_template_docx(template=template_obj, path=path, data=data)
        '''
        if template_obj.template_type == 2:  # 固定模版
            docx_freemarker.make_template_docx(template=template_obj, path=path, data=data)
        else:
            for row in data:
                if row['type'] == 'module':  # 报告名称
                    time = row['params']['time'] if 'time' in row['params'] else ''
                    self.create_docx_modle(content=row['params']['title'], time=time)
                elif row['type'] == 'title1':  # 一级标题
                    self.create_docx_title(content=row['params']['title'], level=1)
                elif row['type'] == 'title2':  # 二级标题
                    self.create_docx_title(content=row['params']['title'], level=2)
                elif row['type'] == 'title3':  # 三级标题
                    self.create_docx_title(content=row['params']['title'], level=3)
                elif row['type'] == 'title4':  # 四级标题
                    self.create_docx_title(content=row['params']['title'], level=4)
                elif row['type'] == 'text':  # 正文
                    self.create_docx_text(content=row['params']['title'])
                elif row['type'] == 'echart':  # 图表
                    self.create_docx_chart(row=row, path=path)
                else:
                    logger.info(row['type'])
            self.docx.save(path)
        '''

    def create_docx_modle(self, content, time):
        '''
        报告名称
        '''
        if content:
            if time:
                title = '{}({})'.format(content, time)
            else:
                title = content
            self.docx.add_heading(title, level=0)

    def create_docx_title(self, content, level):
        '''
        标题
        '''
        self.docx.add_heading(content, level=level)

    def create_docx_text(self, content):
        '''
        正文
        '''
        paragraph = self.docx.add_paragraph()
        run = paragraph.add_run(text=content)
        run.font.name = u'宋体'
        run.font.size = Pt(11)

    def create_chart_number(self, name, data):
        '''
        计数类型
        '''
        count = 0
        for item in data['data']:
            for _ in item['data']:
                count += _
        paragraph = self.docx.add_paragraph()
        run = paragraph.add_run(text='{}:{}'.format(name, str(count)))
        run.font.name = u'宋体'
        run.font.size = Pt(10)

    def create_docx_chart(self, row, path):
        if row['chart_type'] == 'table':  # 表格类型
            self.add_table(name=row['name'], data=row['data'], data_type=row['data_type'])
        elif row['chart_type'] == 'table_merge':  # 合并表格类型
            data = copy.deepcopy(row)
            self.add_table_merge(path=path, data=data['data'], merges=data['merges'], widths=data['widths'])
        elif row['chart_type'] == 'number':  # 计数
            self.create_chart_number(name=row['name'], data=row['data'])
        else:
            if 'has_table' in row and row['has_table']:
                table_data = row['data']
            else:
                table_data = None
            if 'data_type' in row and row['data_type']:
                data_type = row['data_type']
            else:
                data_type = None

            self.generate_section(title=row['name'],
                                  desc=row['description'],
                                  remark=row['remark'],
                                  chart=row['chart_path'],
                                  table_data=table_data,
                                  data_type=data_type)
